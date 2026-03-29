#!/usr/bin/env python3
"""Generate a legal demand letter for debt collection - דלוברי בע"מ vs ABC לוגיסטיקה בע"מ"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_rtl_paragraph(paragraph):
    """Set paragraph direction to RTL."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_rtl_run(run):
    """Set run direction to RTL."""
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def add_rtl_paragraph(doc, text="", style=None, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                      font_name="David", underline=False, italic=False, color=None, space_after=Pt(6), space_before=Pt(0)):
    """Add an RTL paragraph with formatting."""
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.5

    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        run.underline = underline
        run.italic = italic
        set_rtl_run(run)
        if color:
            run.font.color.rgb = color
        # Set complex script font (for Hebrew)
        rPr = run._element.get_or_add_rPr()
        cs_font = OxmlElement('w:rFonts')
        cs_font.set(qn('w:cs'), font_name)
        rPr.append(cs_font)
        cs_size = OxmlElement('w:szCs')
        cs_size.set(qn('w:val'), str(size * 2))
        rPr.append(cs_size)
        if bold:
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)

    return p

def add_mixed_run(paragraph, text, bold=False, underline=False, italic=False, font_name="David", size=12):
    """Add a run with formatting to an existing paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    set_rtl_run(run)
    rPr = run._element.get_or_add_rPr()
    cs_font = OxmlElement('w:rFonts')
    cs_font.set(qn('w:cs'), font_name)
    rPr.append(cs_font)
    cs_size = OxmlElement('w:szCs')
    cs_size.set(qn('w:val'), str(size * 2))
    rPr.append(cs_size)
    if bold:
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)
    return run

def setup_header_footer(doc):
    """Set up header and footer."""
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    set_rtl_paragraph(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header_lines = [
        "משרד עורכי דין דלוברי",
        "רח' רוטשילד 45, תל אביב 6578401",
        "טל: 03-1234567 | פקס: 03-1234568 | דוא\"ל: office@delawvery.co.il"
    ]
    for i, line in enumerate(header_lines):
        if i > 0:
            run = hp.add_run("\n")
        run = hp.add_run(line)
        run.font.name = "David"
        run.font.size = Pt(9) if i > 0 else Pt(11)
        run.bold = (i == 0)
        set_rtl_run(run)
        rPr = run._element.get_or_add_rPr()
        cs_font = OxmlElement('w:rFonts')
        cs_font.set(qn('w:cs'), "David")
        rPr.append(cs_font)
        cs_size = OxmlElement('w:szCs')
        cs_size.set(qn('w:val'), str((9 if i > 0 else 11) * 2))
        rPr.append(cs_size)
        if i == 0:
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)

    # Add line under header
    pPr = hp._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    set_rtl_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add line above footer
    fpPr = fp._element.get_or_add_pPr()
    fpBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '000000')
    fpBdr.append(top)
    fpPr.append(fpBdr)

    run = fp.add_run("מסמך חסוי — יחסי עו\"ד-לקוח")
    run.font.name = "David"
    run.font.size = Pt(8)
    run.italic = True
    set_rtl_run(run)
    rPr = run._element.get_or_add_rPr()
    cs_font = OxmlElement('w:rFonts')
    cs_font.set(qn('w:cs'), "David")
    rPr.append(cs_font)


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default RTL for entire document
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        sectPr = OxmlElement('w:sectPr')
        body.append(sectPr)

    setup_header_footer(doc)

    # === Document content ===

    # Date
    add_rtl_paragraph(doc, "ט\"ו באדר תשפ\"ו", size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_rtl_paragraph(doc, "15 במרץ 2026", size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                      space_after=Pt(12))

    # Recipient
    add_rtl_paragraph(doc, "לכבוד", size=12, bold=True)
    add_rtl_paragraph(doc, "חברת ABC לוגיסטיקה בע\"מ", size=12, bold=True)
    add_rtl_paragraph(doc, "ח.פ. 515000000", size=12)
    add_rtl_paragraph(doc, "", space_after=Pt(6))

    # Sending method
    add_rtl_paragraph(doc, "נשלח בדואר רשום ובדוא\"ל", size=12, underline=True,
                      space_after=Pt(12))

    # Subject line
    p = add_rtl_paragraph(doc, "", size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=Pt(6), space_before=Pt(6))
    add_mixed_run(p, "הנדון: מכתב התראה ודרישה לתשלום חוב בסך ", bold=True, size=14)
    add_mixed_run(p, "85,000 ₪ בתוספת מע\"מ", bold=True, underline=True, size=14)

    # Greeting
    add_rtl_paragraph(doc, "א.ג.נ.,", size=12, space_before=Pt(12), space_after=Pt(12))

    # Section 1 - Introduction
    add_rtl_paragraph(doc, "1. מבוא", size=14, bold=True, space_before=Pt(12))

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "הננו פונים אליכם בשם מרשתנו, ")
    add_mixed_run(p, "דלוברי בע\"מ, ח.פ. 516840319", bold=True)
    add_mixed_run(p, " (להלן: \"")
    add_mixed_run(p, "מרשתנו", bold=True)
    add_mixed_run(p, "\"), בדרישה נחרצת וחד-משמעית לתשלום מלוא החוב המגיע לה מאת חברתכם, ")
    add_mixed_run(p, "חברת ABC לוגיסטיקה בע\"מ, ח.פ. 515000000", bold=True)
    add_mixed_run(p, " (להלן: \"")
    add_mixed_run(p, "החייבת", bold=True)
    add_mixed_run(p, "\"), וכמפורט להלן.")

    # Section 2 - Facts
    add_rtl_paragraph(doc, "2. העובדות", size=14, bold=True, space_before=Pt(12))

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "2.1  ")
    add_mixed_run(p, "ביום 1.9.2025", bold=True)
    add_mixed_run(p, " נכרת הסכם שירותי שליחויות בין מרשתנו לבין החייבת (להלן: \"")
    add_mixed_run(p, "ההסכם", bold=True)
    add_mixed_run(p, "\"), במסגרתו התחייבה מרשתנו לספק לחייבת שירותי שליחויות משפטיות.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "2.2  מרשתנו קיימה את כלל התחייבויותיה על-פי ההסכם במלואן ובמועדן, וסיפקה לחייבת את שירותי השליחויות כסדרם ובאיכות מקצועית גבוהה.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "2.3  בתמורה לשירותים האמורים, התחייבה החייבת לשלם למרשתנו סך של ")
    add_mixed_run(p, "85,000 ₪ (שמונים וחמישה אלף שקלים חדשים) בתוספת מע\"מ כדין", bold=True)
    add_mixed_run(p, ".")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "2.4  מועד התשלום על-פי ההסכם נקבע ליום ")
    add_mixed_run(p, "1.12.2025", bold=True)
    add_mixed_run(p, ". למרבה הצער, חרף פניות חוזרות ונשנות מצד מרשתנו, החייבת טרם שילמה את החוב, כולו או חלקו.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "2.5  נכון למועד מכתב זה, חלפו למעלה מ-")
    add_mixed_run(p, "שלושה חודשים", bold=True)
    add_mixed_run(p, " ממועד התשלום שנקבע בהסכם, והחוב טרם סולק.")

    # Section 3 - Legal basis
    add_rtl_paragraph(doc, "3. הבסיס המשפטי", size=14, bold=True, space_before=Pt(12))

    # 3.1 - Contract Remedies
    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "3.1  על-פי ")
    add_mixed_run(p, "חוק החוזים (תרופות בשל הפרת חוזה), התשל\"א-1970", bold=True, underline=True)
    add_mixed_run(p, ":")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      א.  ")
    add_mixed_run(p, "סעיף 2", bold=True)
    add_mixed_run(p, " — הפרת חוזה היא מעשה או מחדל שהם בניגוד לחוזה. אי-תשלום החוב במועד מהווה הפרה של ההסכם.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      ב.  ")
    add_mixed_run(p, "סעיף 7(א)", bold=True)
    add_mixed_run(p, " — הנפגע זכאי לאכיפת החוזה. מרשתנו זכאית לדרוש את קיום ההתחייבות לתשלום במלואה.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      ג.  ")
    add_mixed_run(p, "סעיף 10", bold=True)
    add_mixed_run(p, " — הנפגע זכאי לפיצויים בעד הנזק שנגרם לו עקב ההפרה. מרשתנו שומרת על זכותה לדרוש פיצויים בגין כל נזק שנגרם לה.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      ד.  ")
    add_mixed_run(p, "סעיף 11(א) ו-(ב)", bold=True)
    add_mixed_run(p, " — מרשתנו זכאית לפיצויים ללא הוכחת נזק (פיצויים מוסכמים), ככל שנקבעו בהסכם, וכן לפיצויים בגין נזק ממשי שנגרם לה.")

    # 3.2 - General Contract Law
    p = add_rtl_paragraph(doc, "", size=12, space_before=Pt(6))
    add_mixed_run(p, "3.2  על-פי ")
    add_mixed_run(p, "חוק החוזים (חלק כללי), התשל\"ג-1973", bold=True, underline=True)
    add_mixed_run(p, ":")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      א.  ")
    add_mixed_run(p, "סעיף 39", bold=True)
    add_mixed_run(p, " — חובת תום-לב בקיום חוזה. התנהלות החייבת, שאינה משלמת את חובה חרף פניות חוזרות, מהווה חוסר תום-לב מובהק.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      ב.  ")
    add_mixed_run(p, "סעיף 61(ב)", bold=True)
    add_mixed_run(p, " — הוראות חוק זה, לרבות חובת תום-הלב, חלות גם מחוץ לדיני החוזים.")

    # 3.3 - Bills of Exchange
    p = add_rtl_paragraph(doc, "", size=12, space_before=Pt(6))
    add_mixed_run(p, "3.3  על-פי ")
    add_mixed_run(p, "פקודת השטרות [נוסח חדש]", bold=True, underline=True)
    add_mixed_run(p, ":")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      ככל שניתנו שטרות או שיקים כבטוחה או כאמצעי תשלום במסגרת ההסכם ולא כובדו — מרשתנו שומרת על מלוא זכויותיה לפעול בהתאם לפקודת השטרות, לרבות הגשת שיקים לביצוע בלשכת ההוצאה לפועל.")

    # Section 4 - Interest and linkage
    add_rtl_paragraph(doc, "4. ריבית והצמדה", size=14, bold=True, space_before=Pt(12))

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "4.1  מרשתנו שומרת על זכותה לדרוש ריבית פיגורים והפרשי הצמדה על סכום החוב מיום ")
    add_mixed_run(p, "1.12.2025", bold=True)
    add_mixed_run(p, " (מועד התשלום החוזי) ועד למועד התשלום בפועל, בהתאם ל")
    add_mixed_run(p, "חוק פסיקת ריבית והצמדה, התשכ\"א-1961", bold=True, underline=True)
    add_mixed_run(p, ".")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "4.2  כמו כן, מרשתנו שומרת על זכותה לדרוש ריבית והצמדה בהתאם לתנאי ההסכם, ככל שנקבעו בו שיעורי ריבית מוסכמים.")

    # Section 5 - Demand
    add_rtl_paragraph(doc, "5. הדרישה", size=14, bold=True, space_before=Pt(12))

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "לאור כל האמור לעיל, הנכם נדרשים בזאת לשלם למרשתנו את מלוא סכום החוב בסך של ")
    add_mixed_run(p, "85,000 ₪ (שמונים וחמישה אלף שקלים חדשים) בתוספת מע\"מ כדין", bold=True)
    add_mixed_run(p, ", בתוספת ריבית פיגורים והפרשי הצמדה כאמור לעיל, וזאת ")
    add_mixed_run(p, "בתוך 14 ימים ממועד קבלת מכתב זה", bold=True, underline=True)
    add_mixed_run(p, ".")

    # Section 6 - Warning
    add_rtl_paragraph(doc, "6. אזהרה על הליכים משפטיים", size=14, bold=True, space_before=Pt(12))

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "6.1  מכתב זה מהווה התראה אחרונה בטרם נקיטת הליכים משפטיים.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "6.2  היה ולא תפעלו לסילוק מלוא החוב במועד שנקבע לעיל, תאלץ מרשתנו לפעול בכל האמצעים העומדים לרשותה, ובכלל זה:")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      א.  הגשת תביעה אזרחית לבית המשפט המוסמך לתשלום מלוא החוב, בתוספת ריבית, הפרשי הצמדה, הוצאות משפט ושכר טרחת עו\"ד;")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      ב.  נקיטת הליכי הוצאה לפועל, לרבות הטלת עיקולים על חשבונות בנק, נכסים וזכויות של החייבת;")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "      ג.  נקיטת כל הליך משפטי נוסף שמרשתנו תמצא לנכון, לרבות פנייה לרשם החברות.")

    p = add_rtl_paragraph(doc, "", size=12, space_before=Pt(6))
    add_mixed_run(p, "6.3  מובהר כי במקרה של נקיטת הליכים משפטיים, ")
    add_mixed_run(p, "ייתבעו מלוא ההוצאות ושכר טרחת עו\"ד מהחייבת", bold=True)
    add_mixed_run(p, ".")

    # Section 7 - Reservation of rights
    add_rtl_paragraph(doc, "7. שמירת זכויות", size=14, bold=True, space_before=Pt(12))

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "7.1  אין באמור במכתב זה כדי למצות ו/או לגרוע מזכויותיה ו/או מטענותיה של מרשתנו על-פי כל דין ו/או הסכם.")

    p = add_rtl_paragraph(doc, "", size=12)
    add_mixed_run(p, "7.2  מכתב זה ישמש כראיה בהליכים משפטיים, ככל שיידרשו.")

    # Closing
    add_rtl_paragraph(doc, "", space_after=Pt(12))

    add_rtl_paragraph(doc, "בכבוד רב,", size=12, space_before=Pt(12))
    add_rtl_paragraph(doc, "", space_after=Pt(24))
    add_rtl_paragraph(doc, "___________________________", size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_rtl_paragraph(doc, "עו\"ד [שם עורך הדין]", size=12, bold=True)
    add_rtl_paragraph(doc, "מ.ר. [מספר רישיון]", size=12)
    add_rtl_paragraph(doc, "בא/ת כוח דלוברי בע\"מ", size=12)

    # CC
    add_rtl_paragraph(doc, "", space_after=Pt(12))
    add_rtl_paragraph(doc, "העתקים:", size=10, underline=True)
    add_rtl_paragraph(doc, "דלוברי בע\"מ — מרשתנו", size=10)
    add_rtl_paragraph(doc, "התיק", size=10)

    # Save
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "מכתב_התראה_גביית_חוב_ABC_לוגיסטיקה.docx")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    main()
