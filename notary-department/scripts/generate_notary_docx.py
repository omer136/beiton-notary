#!/usr/bin/env python3
"""
מחולל DOCX נוטריוני — מייצר אישור נכונות תרגום בפורמט רשמי.

שימוש כמודול:
    from generate_notary_docx import generate_translation_docx
    path = generate_translation_docx(
        original_text="...",
        translated_text="...",
        source_lang="he", target_lang="en",
        document_type="תעודת גירושין",
        client_name="ישראל ישראלי",
        word_count=300,
    )

שימוש ב-CLI:
    python generate_notary_docx.py --original-file source.txt --translation-file trans.txt \
        --source-lang he --target-lang en --doc-type "תעודת גירושין" \
        --client "ישראל ישראלי" --words 300

    python generate_notary_docx.py --demo   # הרצת דוגמה מובנית
"""

from __future__ import annotations

import argparse
import json
import math
import os
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths & Config
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent.parent
CONFIG_DIR = BASE_DIR / "config"
OUTPUT_DIR = BASE_DIR / "templates" / "translation_docx"
PRICING_FILE = CONFIG_DIR / "notary_pricing_2026.json"

# Notary details (defaults — override via params)
DEFAULT_NOTARY = {
    "name_he": "עו״ד דניאל בית-און",
    "name_en": "Adv. Daniel Beit-On",
    "license_no": "54321",
    "address_he": "רח׳ רוטשילד 42, תל אביב-יפו 6688312",
    "address_en": "42 Rothschild Blvd., Tel Aviv-Yafo 6688312",
    "phone": "03-1234567",
    "fax": "03-1234568",
    "email": "notary@firm.co.il",
}

FIRM = {
    "name_he": "משרד עורכי דין בית-און ושות׳",
    "name_en": "Beit-On & Co. Law Offices",
}

# Running certificate number (in production: read from DB/file and increment)
CERT_COUNTER_FILE = CONFIG_DIR / ".cert_counter"

LANG_NAMES = {
    "he": ("עברית", "Hebrew"),
    "en": ("אנגלית", "English"),
    "ru": ("רוסית", "Russian"),
    "ar": ("ערבית", "Arabic"),
    "fr": ("צרפתית", "French"),
    "de": ("גרמנית", "German"),
    "es": ("ספרדית", "Spanish"),
}

# ---------------------------------------------------------------------------
# Pricing
# ---------------------------------------------------------------------------

def calc_notary_fee(word_count: int, notary_translates: bool = True) -> dict:
    with open(PRICING_FILE, encoding="utf-8") as f:
        pricing = json.load(f)
    vat_rate = pricing["_meta"]["vat_rate"]
    items = pricing["pricing"]["translation_approval"]["items"]

    fee = items[0]["amount"]  # first 100
    remaining = max(0, word_count - 100)
    blocks = min(math.ceil(remaining / 100), 9) if remaining > 0 else 0
    fee += blocks * items[1]["amount"]
    above_1000 = max(0, word_count - 1000)
    if above_1000 > 0:
        fee += math.ceil(above_1000 / 100) * items[2]["amount"]

    surcharge = 0
    if notary_translates:
        surcharge = fee * pricing["pricing"]["translation_approval"]["translation_by_notary_surcharge"]["surcharge_rate"]

    total_fee = fee + surcharge
    vat = round(total_fee * vat_rate, 2)
    return {
        "base_fee": fee,
        "surcharge": surcharge,
        "total_before_vat": total_fee,
        "vat_rate": vat_rate,
        "vat": vat,
        "total": round(total_fee + vat, 2),
        "regulation_ref": "פרט 3",
    }

# ---------------------------------------------------------------------------
# Certificate number
# ---------------------------------------------------------------------------

def get_next_cert_number() -> str:
    year = date.today().year
    counter = 1
    if CERT_COUNTER_FILE.exists():
        try:
            data = json.loads(CERT_COUNTER_FILE.read_text())
            if data.get("year") == year:
                counter = data["counter"] + 1
            # else: new year, reset
        except (json.JSONDecodeError, KeyError):
            pass
    CERT_COUNTER_FILE.write_text(json.dumps({"year": year, "counter": counter}))
    return f"{counter:04d}/{year}"

# ---------------------------------------------------------------------------
# Helpers: font & paragraph
# ---------------------------------------------------------------------------

FONT_HE = "David"
FONT_EN = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5
MARGIN = Cm(2.5)

DARK_BLUE = RGBColor(0x1A, 0x3C, 0x6E)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, color_hex: str):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _make_rtl_pPr(pPr):
    """Add <w:bidi/> and <w:jc w:val="right"/> to paragraph properties."""
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))
    # Remove any existing jc first, then set right
    for existing_jc in pPr.findall(qn('w:jc')):
        pPr.remove(existing_jc)
    pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="right"/>'))


def _make_rtl_run(run, font_name: str = FONT_HE, font_size_pt: int = 24):
    """Set full RTL run properties: rtl flag, rFonts with all slots, szCs.
    Must be called INSTEAD of setting run.font.name (not in addition)."""
    rPr = run._element.get_or_add_rPr()
    # Remove any existing rFonts (python-docx may have added one)
    for existing in rPr.findall(qn('w:rFonts')):
        rPr.remove(existing)
    # <w:rtl/>
    rPr.append(parse_xml(f'<w:rtl {nsdecls("w")}/>'))
    # <w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David" w:hint="cs"/>
    rPr.append(parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:cs="{font_name}" w:hint="cs"/>'
    ))
    # Remove any existing szCs
    for existing in rPr.findall(qn('w:szCs')):
        rPr.remove(existing)
    # <w:szCs w:val="24"/> (size in half-points)
    rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{font_size_pt}"/>'))


def _make_ltr_pPr(pPr):
    """Explicitly set paragraph to LTR: <w:bidi w:val="0"/> + <w:jc w:val="left"/>."""
    # Remove any existing bidi
    for existing in pPr.findall(qn('w:bidi')):
        pPr.remove(existing)
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="0"/>'))
    # Remove any existing jc, then set left
    for existing in pPr.findall(qn('w:jc')):
        pPr.remove(existing)
    pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="left"/>'))


def _make_ltr_run(run, font_name: str = FONT_EN):
    """Set explicit LTR run properties: no rtl, proper rFonts."""
    rPr = run._element.get_or_add_rPr()
    # Remove any existing rFonts
    for existing in rPr.findall(qn('w:rFonts')):
        rPr.remove(existing)
    # Remove any rtl element
    for existing in rPr.findall(qn('w:rtl')):
        rPr.remove(existing)
    rPr.append(parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:cs="{font_name}"/>'
    ))


def set_document_rtl(doc):
    """Set the entire document default direction to RTL."""
    # Set bidi on the body-level sectPr
    for section in doc.sections:
        sectPr = section._sectPr
        sectPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))

    # Set default paragraph bidi in document defaults (styles.xml)
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = parse_xml(f'<w:docDefaults {nsdecls("w")}/>')
        styles_element.insert(0, docDefaults)
    pPrDefault = docDefaults.find(qn('w:pPrDefault'))
    if pPrDefault is None:
        pPrDefault = parse_xml(f'<w:pPrDefault {nsdecls("w")}/>')
        docDefaults.append(pPrDefault)
    pPr = pPrDefault.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        pPrDefault.append(pPr)
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))


def add_paragraph(doc, text: str, font_name: str = FONT_EN, font_size=FONT_SIZE,
                  bold: bool = False, color: RGBColor = BLACK,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after: Pt = Pt(6),
                  space_before: Pt = Pt(0), is_rtl: bool = False):
    """Add a styled paragraph with proper RTL/LTR handling."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = LINE_SPACING

    pPr = p._element.get_or_add_pPr()

    if is_rtl:
        _make_rtl_pPr(pPr)
    else:
        _make_ltr_pPr(pPr)

    run = p.add_run(text)
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = color

    if is_rtl:
        _make_rtl_run(run, font_name, int(font_size.pt * 2) if hasattr(font_size, 'pt') else 24)
    else:
        _make_ltr_run(run, font_name)

    return p


def add_bilingual(doc, text_he: str, text_en: str,
                  bold: bool = False, size=FONT_SIZE,
                  color: RGBColor = BLACK):
    """Add two separate paragraphs: Hebrew RTL right-aligned, then English LTR left-aligned."""
    size_half_pt = size.pt * 2 if hasattr(size, 'pt') else 24

    # Hebrew paragraph — RTL, right-aligned
    p_he = doc.add_paragraph()
    p_he.paragraph_format.space_after = Pt(1)
    p_he.paragraph_format.line_spacing = LINE_SPACING
    pPr_he = p_he._element.get_or_add_pPr()
    _make_rtl_pPr(pPr_he)

    run_he = p_he.add_run(text_he)
    run_he.font.size = size
    run_he.font.bold = bold
    run_he.font.color.rgb = color
    _make_rtl_run(run_he, FONT_HE, int(size_half_pt))

    # English paragraph — explicit LTR, left-aligned
    p_en = doc.add_paragraph()
    p_en.paragraph_format.space_after = Pt(4)
    p_en.paragraph_format.line_spacing = LINE_SPACING
    pPr_en = p_en._element.get_or_add_pPr()
    _make_ltr_pPr(pPr_en)

    run_en = p_en.add_run(text_en)
    run_en.font.size = size
    run_en.font.bold = bold
    run_en.font.color.rgb = color
    _make_ltr_run(run_en, FONT_EN)

    return p_he, p_en


def add_separator(doc, char: str = "─", width: int = 55):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Explicit LTR + center for separator
    pPr = p._element.get_or_add_pPr()
    for existing in pPr.findall(qn('w:bidi')):
        pPr.remove(existing)
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="0"/>'))
    for existing in pPr.findall(qn('w:jc')):
        pPr.remove(existing)
    pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
    run = p.add_run(char * width)
    run.font.color.rgb = GRAY
    run.font.size = Pt(10)
    _make_ltr_run(run, FONT_EN)


def _style_header_footer_run(run, font_name, font_size_pt, color, is_he=False):
    """Apply consistent styling to header/footer runs."""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = color
    if is_he:
        _make_rtl_run(run, FONT_HE, font_size_pt * 2)
    else:
        _make_ltr_run(run, font_name)


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def generate_translation_docx(
    original_text: str,
    translated_text: str,
    source_lang: str = "he",
    target_lang: str = "en",
    document_type: str = "תעודה",
    document_type_en: str = "Certificate",
    client_name: str = "",
    word_count: int | None = None,
    page_count: int = 1,
    notary_translates: bool = True,
    notary: dict | None = None,
    cert_number: str | None = None,
    output_path: str | None = None,
) -> Path:
    """Generate a notarial translation certificate DOCX.

    Returns the path to the generated file.
    """
    notary = notary or DEFAULT_NOTARY
    if word_count is None:
        word_count = len(original_text.split())
    if cert_number is None:
        cert_number = get_next_cert_number()

    fees = calc_notary_fee(word_count, notary_translates)
    today = date.today()
    src_name_he, src_name_en = LANG_NAMES.get(source_lang, (source_lang, source_lang))
    tgt_name_he, tgt_name_en = LANG_NAMES.get(target_lang, (target_lang, target_lang))

    # --- Create document ---
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # Set document-level RTL
    set_document_rtl(doc)

    # --- HEADER ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    hp_pPr = hp._element.get_or_add_pPr()
    for b in hp_pPr.findall(qn('w:bidi')):
        hp_pPr.remove(b)
    hp_pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="0"/>'))
    for j in hp_pPr.findall(qn('w:jc')):
        hp_pPr.remove(j)
    hp_pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))

    run_firm = hp.add_run(FIRM["name_he"] + "  |  " + FIRM["name_en"])
    run_firm.font.bold = True
    _style_header_footer_run(run_firm, FONT_EN, 10, DARK_BLUE, is_he=True)

    hp2 = header.add_paragraph()
    hp2.paragraph_format.space_after = Pt(0)
    hp2_pPr = hp2._element.get_or_add_pPr()
    for b in hp2_pPr.findall(qn('w:bidi')):
        hp2_pPr.remove(b)
    hp2_pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="0"/>'))
    for j in hp2_pPr.findall(qn('w:jc')):
        hp2_pPr.remove(j)
    hp2_pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))

    run_addr = hp2.add_run(
        f"{notary['address_he']}  |  ☎ {notary['phone']}  |  📠 {notary['fax']}  |  ✉ {notary['email']}"
    )
    _style_header_footer_run(run_addr, FONT_EN, 8, GRAY, is_he=True)

    # Header bottom border
    hp2_pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    ))

    # --- FOOTER ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(6)
    fp_pPr = fp._element.get_or_add_pPr()
    for b in fp_pPr.findall(qn('w:bidi')):
        fp_pPr.remove(b)
    fp_pPr.append(parse_xml(f'<w:bidi {nsdecls("w")} w:val="0"/>'))
    for j in fp_pPr.findall(qn('w:jc')):
        fp_pPr.remove(j)
    fp_pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))

    fp_pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    ))

    run_f = fp.add_run("מסמך חסוי — יחסי עו״ד-לקוח  |  Confidential — Attorney-Client Privilege")
    _style_header_footer_run(run_f, FONT_EN, 8, GRAY, is_he=True)

    # ===================================================================
    # BODY
    # ===================================================================

    # --- Title ---
    add_bilingual(
        doc,
        f"אישור נכונות תרגום מס׳ {cert_number}",
        f"Certificate of Translation Accuracy No. {cert_number}",
        bold=True, size=Pt(16), color=DARK_BLUE,
    )

    add_separator(doc, "═")

    # --- Section 1: Notary Details ---
    add_bilingual(
        doc,
        "חלק א׳ — פרטי הנוטריון",
        "Part A — Notary Details",
        bold=True, size=Pt(13), color=DARK_BLUE,
    )

    add_paragraph(
        doc,
        f"אני, {notary['name_he']}, נוטריון מס׳ {notary['license_no']}, "
        f"מאשר בזאת כי:",
        font_name=FONT_HE, is_rtl=True,
    )
    add_paragraph(
        doc,
        f"I, {notary['name_en']}, Notary Public No. {notary['license_no']}, "
        f"hereby certify that:",
        font_name=FONT_EN, alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )

    add_separator(doc)

    # --- Section 2: Document Description ---
    add_bilingual(
        doc,
        "חלק ב׳ — תיאור המסמך המקורי",
        "Part B — Original Document Description",
        bold=True, size=Pt(13), color=DARK_BLUE,
    )

    # Description table — RTL: value on right (col 0), label on left (col 1)
    desc_table = doc.add_table(rows=5, cols=2)
    desc_table.style = "Table Grid"
    desc_table.autofit = True

    # Set table bidi for RTL layout
    tblPr = desc_table._tbl.tblPr
    tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))

    desc_data = [
        ("סוג המסמך / Document Type", f"{document_type}  |  {document_type_en}"),
        ("שפת מקור / Source Language", f"{src_name_he}  |  {src_name_en}"),
        ("שפת יעד / Target Language", f"{tgt_name_he}  |  {tgt_name_en}"),
        ("מספר עמודים / Page Count", str(page_count)),
        ("מספר מילים / Word Count", str(word_count)),
    ]

    for i, (label, value) in enumerate(desc_data):
        # In RTL table: cell 0 is right (label), cell 1 is left (value)
        cell_label = desc_table.cell(i, 0)
        cell_value = desc_table.cell(i, 1)
        set_cell_shading(cell_label, "F0F4FA")

        p_label = cell_label.paragraphs[0]
        pPr_l = p_label._element.get_or_add_pPr()
        _make_rtl_pPr(pPr_l)
        run_l = p_label.add_run(label)
        run_l.font.size = Pt(10)
        run_l.font.bold = True
        run_l.font.color.rgb = DARK_BLUE
        _make_rtl_run(run_l, FONT_HE, 20)

        p_value = cell_value.paragraphs[0]
        pPr_v = p_value._element.get_or_add_pPr()
        _make_rtl_pPr(pPr_v)
        run_v = p_value.add_run(value)
        run_v.font.size = Pt(11)
        _make_rtl_run(run_v, FONT_HE, 22)

    for row in desc_table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()  # spacer

    add_separator(doc)

    # --- Section 3: Original Text ---
    add_bilingual(
        doc,
        "חלק ג׳ — המסמך המקורי",
        "Part C — Original Document",
        bold=True, size=Pt(13), color=DARK_BLUE,
    )

    is_source_rtl = source_lang in ("he", "ar")
    source_font = FONT_HE if source_lang in ("he",) else FONT_EN

    for para_text in original_text.strip().split("\n"):
        if para_text.strip():
            add_paragraph(
                doc, para_text.strip(),
                font_name=source_font,
                is_rtl=is_source_rtl,
                space_after=Pt(4),
            )

    add_paragraph(
        doc, "— סוף המקור / End of Original —",
        font_name=FONT_EN, color=GRAY, font_size=Pt(10),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True, space_before=Pt(8),
    )

    add_separator(doc)

    # --- Section 4: Translation ---
    add_bilingual(
        doc,
        "חלק ד׳ — התרגום",
        "Part D — Translation",
        bold=True, size=Pt(13), color=DARK_BLUE,
    )

    is_target_rtl = target_lang in ("he", "ar")
    target_font = FONT_HE if target_lang in ("he",) else FONT_EN

    for para_text in translated_text.strip().split("\n"):
        if para_text.strip():
            add_paragraph(
                doc, para_text.strip(),
                font_name=target_font,
                is_rtl=is_target_rtl,
                alignment=WD_ALIGN_PARAGRAPH.LEFT if not is_target_rtl else WD_ALIGN_PARAGRAPH.RIGHT,
                space_after=Pt(4),
            )

    add_paragraph(
        doc, "— סוף התרגום / End of Translation —",
        font_name=FONT_EN, color=GRAY, font_size=Pt(10),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True, space_before=Pt(8),
    )

    add_separator(doc)

    # --- Section 5: Notary Declaration ---
    add_bilingual(
        doc,
        "חלק ה׳ — הצהרת הנוטריון",
        "Part E — Notary Declaration",
        bold=True, size=Pt(13), color=DARK_BLUE,
    )

    declaration_he = (
        f"אני, הח״מ {notary['name_he']}, נוטריון מס׳ {notary['license_no']}, "
        f"מאשר בזאת כי התרגום דלעיל מ{src_name_he} ל{tgt_name_he} הינו תרגום נאמן ומלא "
        f"של המסמך המקורי שהוצג בפניי. "
        f"התרגום בוצע על ידי ושפות אלו ידועות לי."
    )
    add_paragraph(
        doc, declaration_he,
        font_name=FONT_HE, is_rtl=True,
        space_after=Pt(8),
    )

    declaration_en = (
        f"I, the undersigned, {notary['name_en']}, Notary Public No. {notary['license_no']}, "
        f"hereby certify that the above is a true and faithful translation from "
        f"{src_name_en} to {tgt_name_en} of the original document presented to me. "
        f"The translation was made by me and I am proficient in both languages."
    )
    add_paragraph(
        doc, declaration_en,
        font_name=FONT_EN,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=Pt(8),
    )

    add_separator(doc)

    # --- Section 6: Fee Breakdown ---
    add_bilingual(
        doc,
        "חלק ו׳ — ספירת מילים ושכר נוטריוני",
        "Part F — Word Count & Notarial Fee",
        bold=True, size=Pt(13), color=DARK_BLUE,
    )

    fee_table = doc.add_table(rows=6, cols=2)
    fee_table.style = "Table Grid"
    # RTL table
    fee_tblPr = fee_table._tbl.tblPr
    fee_tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))

    fee_data = [
        ("מספר מילים במקור / Source Word Count", str(word_count)),
        ("שכר אישור תרגום / Translation Certification Fee",
         f"{fees['base_fee']:.2f} ₪"),
        ("תוספת תרגום ע״י נוטריון / Notary Translation Surcharge (50%)",
         f"{fees['surcharge']:.2f} ₪"),
        ("סה״כ לפני מע״מ / Subtotal Before VAT",
         f"{fees['total_before_vat']:.2f} ₪"),
        (f"מע״מ / VAT ({int(fees['vat_rate']*100)}%)",
         f"{fees['vat']:.2f} ₪"),
        ("סה״כ / Total",
         f"{fees['total']:.2f} ₪"),
    ]

    for i, (label, value) in enumerate(fee_data):
        cell_l = fee_table.cell(i, 0)
        cell_v = fee_table.cell(i, 1)
        set_cell_shading(cell_l, "F0F4FA")

        p_l = cell_l.paragraphs[0]
        pPr_fl = p_l._element.get_or_add_pPr()
        _make_rtl_pPr(pPr_fl)
        run_l = p_l.add_run(label)
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = DARK_BLUE
        _make_rtl_run(run_l, FONT_HE, 20)

        p_v = cell_v.paragraphs[0]
        pPr_fv = p_v._element.get_or_add_pPr()
        _make_rtl_pPr(pPr_fv)
        is_total = (i == len(fee_data) - 1)
        run_v = p_v.add_run(value)
        run_v.font.size = Pt(11)
        run_v.font.bold = is_total
        _make_rtl_run(run_v, FONT_HE, 22)
        if is_total:
            set_cell_shading(cell_v, "E8F0E8")

    for row in fee_table.rows:
        row.cells[0].width = Cm(10)
        row.cells[1].width = Cm(6)

    add_paragraph(
        doc,
        f"בהתאם לתקנות הנוטריונים (שכר שירותים), תשל״ט-1978, {fees['regulation_ref']}",
        font_name=FONT_HE, font_size=Pt(9), color=GRAY,
        is_rtl=True, space_before=Pt(4),
    )

    add_separator(doc)

    # --- Section 7: Signature Block ---
    add_bilingual(
        doc,
        "חלק ז׳ — חתימה וחותמת",
        "Part G — Signature & Seal",
        bold=True, size=Pt(13), color=DARK_BLUE,
    )

    doc.add_paragraph()  # spacer

    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.autofit = True
    sig_tblPr = sig_table._tbl.tblPr
    sig_tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))

    sig_data = [
        ("תאריך / Date:", f"__{today.strftime('%d / %m / %Y')}__"),
        ("חתימה / Signature:", "________________________"),
        ("חותמת / Seal:", "[מקום לחותמת נוטריון / Notary Seal]"),
        ("", f"{notary['name_he']}\n{notary['name_en']}\nנוטריון מס׳ / Notary No. {notary['license_no']}"),
    ]

    for i, (label, value) in enumerate(sig_data):
        cell_0 = sig_table.cell(i, 0)
        cell_1 = sig_table.cell(i, 1)

        if label:
            p0 = cell_0.paragraphs[0]
            pPr_s0 = p0._element.get_or_add_pPr()
            _make_rtl_pPr(pPr_s0)
            run_s0 = p0.add_run(label)
            run_s0.font.size = Pt(11)
            _make_rtl_run(run_s0, FONT_HE, 22)

        p1 = cell_1.paragraphs[0]
        pPr_s1 = p1._element.get_or_add_pPr()
        _make_rtl_pPr(pPr_s1)
        run_s1 = p1.add_run(value)
        run_s1.font.size = Pt(10 if i >= 2 else 11)
        _make_rtl_run(run_s1, FONT_HE, 20 if i >= 2 else 22)
        if i == 2:
            run_s1.font.color.rgb = GRAY

    for row in sig_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()  # final spacer

    # --- Save ---
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if output_path:
        out = Path(output_path)
    else:
        safe_type = document_type_en.lower().replace(" ", "_").replace("/", "_")
        filename = (
            f"notary_translation_{safe_type}_"
            f"{source_lang}-{target_lang}_{today.isoformat()}.docx"
        )
        out = OUTPUT_DIR / filename

    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Demo
# ---------------------------------------------------------------------------

DEMO_ORIGINAL = """מדינת ישראל
בית הדין הרבני האזורי תל אביב-יפו

תעודת גירושין

מספר תיק: 1234567/2025

אנו מאשרים בזאת כי:

הבעל: ישראל בן דוד כהן, תעודת זהות מספר 012345678
מרח׳ הרצל 15, תל אביב-יפו

האישה: שרה בת משה לוי (לשעבר כהן), תעודת זהות מספר 987654321
מרח׳ דיזנגוף 42, תל אביב-יפו

התגרשו כדת משה וישראל ביום כ״ה בטבת תשפ״ו, המתאים ליום 15 בינואר 2026,
בבית הדין הרבני האזורי תל אביב-יפו.

הגט נמסר מידי הבעל לידי האישה בפני בית הדין.

מתאריך זה שניהם חופשיים להינשא לכל אדם, בכפוף להוראות כל דין.

ניתן בבית הדין הרבני האזורי תל אביב-יפו
ביום כ״ה בטבת תשפ״ו (15.01.2026)

הרב אברהם יצחקי — אב בית הדין
הרב משה דוידוב — דיין
הרב יעקב פרידמן — דיין

חותמת בית הדין הרבני האזורי תל אביב-יפו

מספר אישור: גט/2026/4567"""


DEMO_TRANSLATION = """State of Israel
Regional Rabbinical Court of Tel Aviv-Yafo

Certificate of Divorce

Case Number: 1234567/2025

We hereby certify that:

The Husband: Israel Ben David Cohen, Identity Card No. 012345678
15 Herzl Street, Tel Aviv-Yafo

The Wife: Sarah Bat Moshe Levy (née Cohen), Identity Card No. 987654321
42 Dizengoff Street, Tel Aviv-Yafo

Were divorced in accordance with the Law of Moses and Israel on the 25th of Tevet 5786,
corresponding to January 15, 2026, at the Regional Rabbinical Court of Tel Aviv-Yafo.

The Get (bill of divorce) was delivered by the Husband to the Wife in the presence
of the Court.

From this date, both parties are free to marry any person, subject to the provisions
of any applicable law.

Given at the Regional Rabbinical Court of Tel Aviv-Yafo
on the 25th of Tevet 5786 (January 15, 2026)

Rabbi Avraham Yitzhaki — Head of the Court (Av Beit Din)
Rabbi Moshe Davidov — Judge (Dayan)
Rabbi Yaakov Friedman — Judge (Dayan)

[Official Seal of the Regional Rabbinical Court of Tel Aviv-Yafo]

Certificate Number: GET/2026/4567"""


def run_demo():
    print("\n" + "=" * 60)
    print("  דמו — הפקת DOCX נוטריוני")
    print("  תעודת גירושין, 300 מילים, עברית → אנגלית")
    print("=" * 60 + "\n")

    word_count = 300
    fees = calc_notary_fee(word_count, notary_translates=True)

    print(f"  מילים: {word_count}")
    print(f"  שכר אישור: {fees['base_fee']:.2f} ₪")
    print(f"  תוספת תרגום נוטריון: {fees['surcharge']:.2f} ₪")
    print(f"  סה״כ לפני מע״מ: {fees['total_before_vat']:.2f} ₪")
    print(f"  מע״מ: {fees['vat']:.2f} ₪")
    print(f"  סה״כ: {fees['total']:.2f} ₪")
    print()

    out_path = generate_translation_docx(
        original_text=DEMO_ORIGINAL,
        translated_text=DEMO_TRANSLATION,
        source_lang="he",
        target_lang="en",
        document_type="תעודת גירושין",
        document_type_en="Divorce Certificate",
        client_name="ישראל כהן",
        word_count=word_count,
        page_count=1,
        notary_translates=True,
    )

    print(f"  ✅ קובץ נוצר: {out_path}")
    print(f"  📄 גודל: {out_path.stat().st_size:,} bytes")
    print()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="מחולל DOCX נוטריוני")
    parser.add_argument("--demo", action="store_true", help="הרץ דוגמה מובנית")
    parser.add_argument("--original-file", help="קובץ טקסט מקורי")
    parser.add_argument("--translation-file", help="קובץ תרגום")
    parser.add_argument("--source-lang", default="he", help="שפת מקור")
    parser.add_argument("--target-lang", default="en", help="שפת יעד")
    parser.add_argument("--doc-type", default="תעודה", help="סוג מסמך בעברית")
    parser.add_argument("--doc-type-en", default="Certificate", help="סוג מסמך באנגלית")
    parser.add_argument("--client", default="", help="שם לקוח")
    parser.add_argument("--words", type=int, help="מספר מילים (אופציונלי)")
    parser.add_argument("--pages", type=int, default=1, help="מספר עמודים")
    parser.add_argument("--output", help="נתיב קובץ פלט")

    args = parser.parse_args()

    if args.demo:
        run_demo()
        return

    if not args.original_file or not args.translation_file:
        parser.error("נדרש --original-file ו---translation-file (או --demo)")

    original = Path(args.original_file).read_text(encoding="utf-8")
    translation = Path(args.translation_file).read_text(encoding="utf-8")

    out = generate_translation_docx(
        original_text=original,
        translated_text=translation,
        source_lang=args.source_lang,
        target_lang=args.target_lang,
        document_type=args.doc_type,
        document_type_en=args.doc_type_en,
        client_name=args.client,
        word_count=args.words,
        page_count=args.pages,
        output_path=args.output,
    )
    print(f"✅ {out}")


if __name__ == "__main__":
    main()
