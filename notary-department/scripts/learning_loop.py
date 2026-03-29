#!/usr/bin/env python3
"""
Learning Loop — השוואת טיוטת AI מול גרסה סופית מאושרת של נוטריון.

1) מחלץ טקסט משני קבצי DOCX
2) משווה ומזהה כל הבדל ברמת מילה
3) מסווג כל תיקון ל-9 קטגוריות
4) מדרג חומרה: קוסמטי / קל / בינוני / מהותי / קריטי
5) כותב ניתוח: למה הסוכן טעה ומה הלקח
6) שומר ל-config/learned_corrections.json (Agent 2 קורא משם)
7) דוּח ל-monday.com בורד למידה 18405974782

שימוש:
    python learning_loop.py --draft draft.docx --final final.docx
    python learning_loop.py --demo
"""

from __future__ import annotations

import argparse
import difflib
import hashlib
import json
import math
import os
import re
import sys
import urllib.request
import urllib.error
from datetime import datetime
from pathlib import Path

from docx import Document

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent.parent
CONFIG_DIR = BASE_DIR / "config"
OUTPUT_DIR = BASE_DIR / "templates" / "translation_docx"
CORRECTIONS_FILE = CONFIG_DIR / "learned_corrections.json"
MONDAY_QA_CONFIG = CONFIG_DIR / "monday_qa_config.json"

# ---------------------------------------------------------------------------
# Categories & Severity
# ---------------------------------------------------------------------------

CATEGORIES = {
    "legal_term": {
        "he": "טרמינולוגיה משפטית",
        "desc": "מונח משפטי תורגם לא נכון או לא בנוסח המקובל",
        "markers": [
            "certif", "hereby", "undersigned", "pursuant", "accordance", "witness",
            "whereof", "provisions", "applicable", "prescribed",
            "מאשר", "בזאת", "הח\"מ", "בהתאם", "נוטריון", "תקנות", "חוק",
        ],
    },
    "name_translit": {
        "he": "שם / תעתיק",
        "desc": "שם פרטי, משפחה או מקום תועתק לא נכון",
        "markers": [
            "Israel", "Moshe", "David", "Sarah", "Cohen", "Kohen", "Levy", "Levi",
            "Friedman", "Davidov", "Yitzhaki", "Yafo", "Jaffa", "Herzl", "Hertzel",
            "Ben", "Bat", "בן", "בת", "כהן", "לוי", "ישראל", "שרה", "משה",
        ],
    },
    "date_number": {
        "he": "תאריך / מספר",
        "desc": "תאריך הומר לא נכון, מספר ת.ז. או אסמכתא שונה",
        "markers_re": [
            r'\d{1,2}[./]\d{1,2}[./]\d{2,4}',
            r'\b\d{6,9}\b',
            r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\b',
            r'\b(טבת|שבט|אדר|ניסן|אייר|סיוון|תמוז|אב|אלול|תשרי|חשוון|כסלו)\b',
            r'\bתש[א-ת]+"[א-ת]\b',
        ],
    },
    "format": {
        "he": "פורמט / עיצוב",
        "desc": "ריווח, פונט, פיסוק, רישיות, עיצוב שאינו משנה משמעות",
    },
    "syntax": {
        "he": "תחביר / דקדוק",
        "desc": "דקדוק, סדר מילים, או ניסוח שאינו טבעי",
    },
    "halachic": {
        "he": "מונח הלכתי / דתי",
        "desc": "מונח דתי, הלכתי או מבית דין רבני תורגם לא נכון",
        "markers": [
            "גט", "כדת משה וישראל", "הלכה", "בית הדין הרבני", "דיין",
            "אב בית הדין", "Get", "bill of divorce", "Law of Moses",
            "Rabbinical", "Dayan", "Av Beit Din", "Jewish Law",
        ],
    },
    "government": {
        "he": "מונח ממשלתי / מוסדי",
        "desc": "שם מוסד, משרד ממשלתי או גוף רשמי תורגם שלא לפי הנוסח הרשמי",
        "markers": [
            "Ministry", "Court", "Authority", "Registrar", "Institute",
            "משרד", "בית משפט", "רשות", "רשם", "לשכת", "המוסד",
        ],
    },
    "meaning": {
        "he": "דיוק משמעות",
        "desc": "הטקסט מתורגם אבל המשמעות שונה או מעורפלת",
    },
    "omission": {
        "he": "השמטה / תוספת",
        "desc": "תוכן שהושמט מהטיוטה או שנוסף שלא היה במקור",
    },
}

SEVERITY_LEVELS = {
    "cosmetic":  {"he": "קוסמטי",  "score": 1, "emoji": "⚪"},
    "minor":     {"he": "קל",      "score": 2, "emoji": "🟢"},
    "moderate":  {"he": "בינוני",   "score": 4, "emoji": "🟡"},
    "major":     {"he": "מהותי",    "score": 7, "emoji": "🟠"},
    "critical":  {"he": "קריטי",   "score": 10, "emoji": "🔴"},
}

# Category → default severity
DEFAULT_SEVERITY = {
    "legal_term":    "major",
    "name_translit": "critical",
    "date_number":   "critical",
    "format":        "cosmetic",
    "syntax":        "minor",
    "halachic":      "major",
    "government":    "major",
    "meaning":       "major",
    "omission":      "major",
}


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_paragraphs(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    paras = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            paras.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        paras.append(t)
    return paras


# ---------------------------------------------------------------------------
# Classification engine
# ---------------------------------------------------------------------------

def _check_markers(text: str, markers: list[str]) -> bool:
    low = text.lower()
    return any(m.lower() in low for m in markers)


def _check_markers_re(text: str, patterns: list[str]) -> bool:
    return any(re.search(p, text, re.IGNORECASE) for p in patterns)


def classify(draft: str, final: str) -> str:
    """Classify the category of a change."""
    combined = draft + " " + final

    # 1. Date/number — if numbers actually differ
    draft_nums = set(re.findall(r'\b\d{3,}\b', draft))
    final_nums = set(re.findall(r'\b\d{3,}\b', final))
    if draft_nums != final_nums:
        return "date_number"
    if _check_markers_re(combined, CATEGORIES["date_number"].get("markers_re", [])):
        date_draft = set(re.findall(r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}', draft, re.I))
        date_final = set(re.findall(r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}', final, re.I))
        if date_draft != date_final:
            return "date_number"

    # 2. Halachic (before legal, more specific)
    if _check_markers(combined, CATEGORIES["halachic"].get("markers", [])):
        return "halachic"

    # 3. Name / transliteration
    if _check_markers(combined, CATEGORIES["name_translit"].get("markers", [])):
        return "name_translit"

    # 4. Government / institution
    if _check_markers(combined, CATEGORIES["government"].get("markers", [])):
        return "government"

    # 5. Legal terminology
    if _check_markers(combined, CATEGORIES["legal_term"].get("markers", [])):
        return "legal_term"

    # 6. Omission / addition
    if not draft.strip():
        return "omission"
    if not final.strip():
        return "omission"

    # 7. Format only (punctuation, whitespace, capitalization)
    d_clean = re.sub(r'[^\w]', '', draft.lower())
    f_clean = re.sub(r'[^\w]', '', final.lower())
    if d_clean == f_clean:
        return "format"

    # 8. Syntax (small edits)
    ratio = difflib.SequenceMatcher(None, draft, final).ratio()
    if ratio > 0.85:
        return "syntax"

    # 9. Meaning (larger semantic difference)
    if ratio < 0.6:
        return "meaning"

    return "syntax"


def assess_severity(category: str, draft: str, final: str) -> str:
    """Determine severity, adjusting the default based on change magnitude."""
    base = DEFAULT_SEVERITY.get(category, "moderate")
    ratio = difflib.SequenceMatcher(None, draft, final).ratio()

    # Upgrade if the change is enormous
    if category in ("syntax", "format") and ratio < 0.5:
        return "major"

    # Downgrade if the change is tiny
    if category in ("name_translit", "government") and ratio > 0.95:
        return "moderate"

    return base


# ---------------------------------------------------------------------------
# Root cause analysis
# ---------------------------------------------------------------------------

def analyze_why(category: str, draft: str, final: str, word_changes: list[dict]) -> dict:
    """Generate a WHY analysis: why the AI made this mistake and what's the lesson."""

    wrong_words = [wc["draft"] for wc in word_changes if wc["draft"]]
    correct_words = [wc["final"] for wc in word_changes if wc["final"]]

    analysis = {"root_cause": "", "lesson": "", "prevention": ""}

    if category == "name_translit":
        analysis["root_cause"] = (
            f"ה-AI השתמש בתעתיק לא תקני לשם עברי. "
            f"'{' '.join(wrong_words)}' במקום '{' '.join(correct_words)}'."
        )
        analysis["lesson"] = (
            "שמות עבריים צריכים תעתיק לפי הנוהג המקובל בישראל, "
            "לא לפי כללי transliteration גנריים."
        )
        analysis["prevention"] = (
            "הוסף ל-patterns: " +
            ", ".join(f'"{w}"→"{c}"' for w, c in zip(wrong_words, correct_words))
        )

    elif category == "halachic":
        analysis["root_cause"] = (
            f"ה-AI לא הכיר את הנוסח המשפטי-הלכתי המקובל. "
            f"תרגם ל-'{draft[:60]}' במקום '{final[:60]}'."
        )
        analysis["lesson"] = (
            "מונחים הלכתיים ומבית דין רבני דורשים תרגום ספציפי מקובל. "
            "למשל: 'כדת משה וישראל' = 'in accordance with the Law of Moses and Israel', "
            "לא 'according to Jewish Law'."
        )
        analysis["prevention"] = "הוסף מונח למילון המונחים ול-patterns."

    elif category == "government":
        analysis["root_cause"] = (
            f"ה-AI תרגם שם מוסד/מקום שלא לפי הנוסח הרשמי. "
            f"'{' '.join(wrong_words)}' במקום '{' '.join(correct_words)}'."
        )
        analysis["lesson"] = (
            "שמות מוסדות ומקומות ישראליים יש להם נוסח אנגלי רשמי קבוע."
        )
        analysis["prevention"] = "הוסף ל-patterns ולמילון מוסדות."

    elif category == "legal_term":
        analysis["root_cause"] = (
            f"מונח משפטי תורגם בנוסח לא מקצועי או לא מקובל."
        )
        analysis["lesson"] = "השתמש תמיד בנוסח הפורמלי המקובל במסמכים נוטריוניים."
        analysis["prevention"] = "הוסף ל-patterns ולמילון מונחים."

    elif category == "date_number":
        analysis["root_cause"] = "תאריך או מספר זיהוי שונה — שגיאה קריטית."
        analysis["lesson"] = "מספרים ותאריכים חייבים העתקה מדויקת ללא שינוי."
        analysis["prevention"] = "הוסף בדיקת regex ספציפית לשדה זה."

    elif category == "omission":
        if not draft.strip():
            analysis["root_cause"] = "הטיוטה השמיטה תוכן שקיים במסמך המקורי."
            analysis["lesson"] = "אסור להשמיט שום שורה, גם אם נראית פורמלית."
        else:
            analysis["root_cause"] = "הטיוטה הוסיפה תוכן שלא קיים במקור."
            analysis["lesson"] = "אסור להוסיף תוכן שלא קיים במסמך המקורי."
        analysis["prevention"] = "בדיקת ספירת משפטים ב-QA."

    elif category == "meaning":
        analysis["root_cause"] = "התרגום שינה את המשמעות המקורית."
        analysis["lesson"] = "תרגום נוטריוני חייב לשמר משמעות מדויקת, לא לפרש."
        analysis["prevention"] = "בדיקת back-translation ב-QA."

    elif category == "syntax":
        analysis["root_cause"] = "ניסוח לא טבעי או שגיאת דקדוק."
        analysis["lesson"] = "המשפט צריך להישמע טבעי בשפת היעד."
        analysis["prevention"] = "בדיקת grammar ב-QA."

    elif category == "format":
        analysis["root_cause"] = "שינוי פורמט — פיסוק, רישיות, ריווח."
        analysis["lesson"] = "מסמכים נוטריוניים דורשים פורמט אחיד."
        analysis["prevention"] = "בדיקת format ב-QA."

    return analysis


# ---------------------------------------------------------------------------
# Comparison engine
# ---------------------------------------------------------------------------

def compare_documents(draft_path: Path, final_path: Path) -> dict:
    draft_paras = extract_paragraphs(draft_path)
    final_paras = extract_paragraphs(final_path)

    corrections = []
    cid = 0

    matcher = difflib.SequenceMatcher(None, draft_paras, final_paras)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue

        if tag == "replace":
            pairs = list(zip(range(i1, i2), range(j1, j2)))
            for di, fi in pairs:
                dp = draft_paras[di]
                fp = final_paras[fi]
                if dp == fp:
                    continue

                # Word-level diff
                dw = dp.split()
                fw = fp.split()
                wm = difflib.SequenceMatcher(None, dw, fw)
                word_changes = []
                for wt, wi1, wi2, wj1, wj2 in wm.get_opcodes():
                    if wt != "equal":
                        word_changes.append({
                            "type": wt,
                            "draft": " ".join(dw[wi1:wi2]),
                            "final": " ".join(fw[wj1:wj2]),
                        })

                cid += 1
                cat = classify(dp, fp)
                sev = assess_severity(cat, dp, fp)
                why = analyze_why(cat, dp, fp, word_changes)

                corrections.append({
                    "id": cid,
                    "op": "replace",
                    "category": cat,
                    "category_he": CATEGORIES[cat]["he"],
                    "severity": sev,
                    "severity_he": SEVERITY_LEVELS[sev]["he"],
                    "draft_text": dp,
                    "final_text": fp,
                    "word_changes": word_changes,
                    "similarity": round(difflib.SequenceMatcher(None, dp, fp).ratio(), 3),
                    "analysis": why,
                })

            # Leftovers
            extra_draft = list(range(len(pairs) + i1, i2))
            extra_final = list(range(len(pairs) + j1, j2))
            for di in extra_draft:
                cid += 1
                corrections.append(_make_omission(cid, "delete", draft_paras[di], ""))
            for fi in extra_final:
                cid += 1
                corrections.append(_make_omission(cid, "insert", "", final_paras[fi]))

        elif tag == "delete":
            for di in range(i1, i2):
                cid += 1
                corrections.append(_make_omission(cid, "delete", draft_paras[di], ""))

        elif tag == "insert":
            for fi in range(j1, j2):
                cid += 1
                corrections.append(_make_omission(cid, "insert", "", final_paras[fi]))

    # Summary
    cat_counts = {}
    sev_counts = {s: 0 for s in SEVERITY_LEVELS}
    total_penalty = 0
    for c in corrections:
        cat_counts[c["category"]] = cat_counts.get(c["category"], 0) + 1
        sev_counts[c["severity"]] += 1
        total_penalty += SEVERITY_LEVELS[c["severity"]]["score"]

    max_penalty = max(len(draft_paras), 1) * 10
    quality = round(max(0, 100 * (1 - total_penalty / max_penalty)), 1)

    overall_sim = round(difflib.SequenceMatcher(
        None, "\n".join(draft_paras), "\n".join(final_paras)
    ).ratio(), 3)

    return {
        "_meta": {
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "draft_file": str(draft_path),
            "final_file": str(final_path),
        },
        "summary": {
            "total_corrections": len(corrections),
            "overall_similarity": overall_sim,
            "quality_score": quality,
            "draft_paragraphs": len(draft_paras),
            "final_paragraphs": len(final_paras),
            "by_category": cat_counts,
            "by_severity": sev_counts,
        },
        "corrections": corrections,
    }


def _make_omission(cid, op, draft, final):
    why = {
        "root_cause": "השמטה מהטיוטה" if op == "insert" else "תוספת בטיוטה שלא קיימת במקור",
        "lesson": "אסור להשמיט או להוסיף תוכן",
        "prevention": "בדיקת ספירת משפטים ב-QA",
    }
    return {
        "id": cid, "op": op, "category": "omission",
        "category_he": "השמטה / תוספת",
        "severity": "major", "severity_he": "מהותי",
        "draft_text": draft, "final_text": final,
        "word_changes": [], "similarity": 0, "analysis": why,
    }


# ---------------------------------------------------------------------------
# Save to learned_corrections.json
# ---------------------------------------------------------------------------

def save_to_learned_corrections(report: dict, report_source: str = ""):
    db = _load_db()

    existing_ids = {r["id"] for r in db["rules"]}
    new_count = 0
    updated_count = 0

    for c in report["corrections"]:
        rule_id = hashlib.sha256(
            f"{c['draft_text']}|||{c['final_text']}".encode()
        ).hexdigest()[:12]

        if rule_id in existing_ids:
            for r in db["rules"]:
                if r["id"] == rule_id:
                    r["occurrences"] += 1
                    r["last_seen"] = datetime.utcnow().isoformat() + "Z"
                    updated_count += 1
                    break
            continue

        # Build word-level patterns
        word_pairs = [
            {"wrong": wc["draft"], "correct": wc["final"]}
            for wc in c.get("word_changes", [])
            if wc["type"] == "replace" and wc["draft"] and wc["final"]
        ]

        db["rules"].append({
            "id": rule_id,
            "category": c["category"],
            "category_he": c["category_he"],
            "severity": c["severity"],
            "severity_he": c["severity_he"],
            "wrong_text": c["draft_text"][:250],
            "correct_text": c["final_text"][:250],
            "word_corrections": word_pairs,
            "analysis": c.get("analysis", {}),
            "source": report_source,
            "learned_at": datetime.utcnow().isoformat() + "Z",
            "occurrences": 1,
        })
        existing_ids.add(rule_id)
        new_count += 1

    # Rebuild quick-lookup patterns
    patterns = {}
    for r in db["rules"]:
        for wp in r.get("word_corrections", []):
            key = wp["wrong"].lower().strip()
            if not key:
                continue
            if key not in patterns:
                patterns[key] = []
            # Deduplicate
            if not any(p["correct"] == wp["correct"] for p in patterns[key]):
                patterns[key].append({
                    "correct": wp["correct"],
                    "category": r["category"],
                    "severity": r["severity"],
                    "rule_id": r["id"],
                })

    db["patterns"] = patterns

    # Statistics
    cat_c = {}
    sev_c = {}
    for r in db["rules"]:
        occ = r.get("occurrences", 1)
        cat_c[r["category"]] = cat_c.get(r["category"], 0) + occ
        sev_c[r["severity"]] = sev_c.get(r["severity"], 0) + occ
    db["statistics"] = {
        "by_category": cat_c,
        "by_severity": sev_c,
        "top_mistakes": sorted(
            [{"id": r["id"],
              "description": _describe_rule(r),
              "occurrences": r.get("occurrences", 1)}
             for r in db["rules"]],
            key=lambda x: -x["occurrences"]
        )[:15],
    }

    db["_meta"]["last_updated"] = datetime.utcnow().isoformat() + "Z"
    db["_meta"]["total_rules"] = len(db["rules"])
    db["_meta"]["total_reports"] += 1

    _save_db(db)

    return {"new": new_count, "updated": updated_count, "total": len(db["rules"])}


def _describe_rule(r):
    pairs = r.get("word_corrections", [])
    if pairs:
        sample = pairs[0]
        return f'{r["category_he"]}: "{sample["wrong"]}" → "{sample["correct"]}"'
    return f'{r["category_he"]}: תיקון'


def _load_db():
    if CORRECTIONS_FILE.exists():
        with open(CORRECTIONS_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {
        "_meta": {
            "description": "תיקונים נלמדים — Agent 2 קורא מכאן לפני כל תרגום",
            "schema_version": 2,
            "created": datetime.utcnow().isoformat() + "Z",
            "last_updated": None,
            "total_rules": 0,
            "total_reports": 0,
        },
        "rules": [],
        "patterns": {},
        "statistics": {},
    }


def _save_db(db):
    with open(CORRECTIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(db, f, ensure_ascii=False, indent=2)


# ---------------------------------------------------------------------------
# monday.com reporting
# ---------------------------------------------------------------------------

def report_to_monday(report: dict, save_result: dict):
    api_key = os.environ.get("MONDAY_API_KEY", "")
    if not api_key:
        print("  ⚠️  MONDAY_API_KEY לא מוגדר — דילוג על monday.com")
        return

    if not MONDAY_QA_CONFIG.exists():
        print("  ⚠️  monday_qa_config.json לא נמצא — דילוג")
        return

    with open(MONDAY_QA_CONFIG, encoding="utf-8") as f:
        qa_cfg = json.load(f)

    board_id = qa_cfg["board_id"]
    cols = qa_cfg["columns"]
    s = report["summary"]

    item_name = (
        f"למידה — {s['total_corrections']} תיקונים, ציון {s['quality_score']}/100 "
        f"({datetime.now().strftime('%d.%m.%Y')})"
    )

    col_values = {}
    col_values[cols["quality_score"]] = str(s["quality_score"])
    col_values[cols["total_corrections"]] = str(s["total_corrections"])
    col_values[cols["similarity"]] = str(round(s["overall_similarity"] * 100, 1))
    col_values[cols["new_rules"]] = str(save_result["new"])
    col_values[cols["report_date"]] = {"date": datetime.now().strftime("%Y-%m-%d")}

    # Status by quality score
    if s["quality_score"] >= 90:
        col_values[cols["status"]] = {"index": qa_cfg["status_labels"]["excellent"]["index"]}
    elif s["quality_score"] >= 70:
        col_values[cols["status"]] = {"index": qa_cfg["status_labels"]["good"]["index"]}
    elif s["quality_score"] >= 50:
        col_values[cols["status"]] = {"index": qa_cfg["status_labels"]["needs_improvement"]["index"]}
    else:
        col_values[cols["status"]] = {"index": qa_cfg["status_labels"]["poor"]["index"]}

    col_values_json = json.dumps(col_values, ensure_ascii=False)

    query = """
    mutation ($boardId: ID!, $name: String!, $cols: JSON!) {
      create_item(board_id: $boardId, item_name: $name, column_values: $cols) { id }
    }
    """
    try:
        result = _monday_call(api_key, query, {
            "boardId": board_id, "name": item_name, "cols": col_values_json,
        })
        item_id = result["data"]["create_item"]["id"]

        # Add detailed update
        body_lines = [f"📊 **Learning Loop Report**\n"]
        body_lines.append(f"• Quality score: **{s['quality_score']}/100**")
        body_lines.append(f"• Similarity: {s['overall_similarity']*100:.1f}%")
        body_lines.append(f"• Corrections: {s['total_corrections']}")
        body_lines.append(f"• New rules: {save_result['new']}, Updated: {save_result['updated']}\n")
        body_lines.append("**By severity:**")
        for sev in ("critical", "major", "moderate", "minor", "cosmetic"):
            cnt = s["by_severity"].get(sev, 0)
            if cnt:
                emoji = SEVERITY_LEVELS[sev]["emoji"]
                body_lines.append(f"  {emoji} {SEVERITY_LEVELS[sev]['he']}: {cnt}")
        body_lines.append("\n**By category:**")
        for cat, cnt in sorted(s["by_category"].items(), key=lambda x: -x[1]):
            body_lines.append(f"  {CATEGORIES[cat]['he']}: {cnt}")
        body_lines.append("\n**Top corrections:**")
        for c in report["corrections"][:5]:
            body_lines.append(f"  {SEVERITY_LEVELS[c['severity']]['emoji']} {c['category_he']}")
            for wc in c.get("word_changes", [])[:2]:
                if wc["draft"] and wc["final"]:
                    body_lines.append(f"    \"{wc['draft']}\" → \"{wc['final']}\"")

        body = "\n".join(body_lines)
        _monday_call(api_key, """
            mutation ($id: ID!, $body: String!) {
              create_update(item_id: $id, body: $body) { id }
            }
        """, {"id": str(item_id), "body": body})

        print(f"  ✅ monday.com: פריט #{item_id} בבורד {board_id}")

    except Exception as e:
        print(f"  ❌ monday.com: {e}")


def _monday_call(api_key, query, variables):
    payload = json.dumps({"query": query, "variables": variables}).encode()
    req = urllib.request.Request(
        "https://api.monday.com/v2", data=payload,
        headers={"Authorization": api_key, "Content-Type": "application/json",
                 "API-Version": "2024-10"},
    )
    with urllib.request.urlopen(req, timeout=30) as resp:
        body = json.loads(resp.read().decode())
    if "errors" in body:
        raise RuntimeError(json.dumps(body["errors"], ensure_ascii=False))
    return body


# ---------------------------------------------------------------------------
# Print
# ---------------------------------------------------------------------------

def print_report(report: dict):
    s = report["summary"]
    print()
    print("═" * 65)
    print("  Learning Loop — דו״ח השוואת טיוטת AI מול גרסה מאושרת")
    print("═" * 65)
    print(f"  ציון איכות:  {s['quality_score']}/100")
    print(f"  דמיון כללי:  {s['overall_similarity']*100:.1f}%")
    print(f"  תיקונים:     {s['total_corrections']}")
    print()

    print("  ── חומרה ──")
    for sev in ("critical", "major", "moderate", "minor", "cosmetic"):
        cnt = s["by_severity"].get(sev, 0)
        if cnt:
            info = SEVERITY_LEVELS[sev]
            bar = "█" * cnt
            print(f"    {info['emoji']} {info['he']:<8} {cnt:>3}  {bar}")

    print()
    print("  ── קטגוריה ──")
    for cat, cnt in sorted(s["by_category"].items(), key=lambda x: -x[1]):
        print(f"    {CATEGORIES[cat]['he']:<22} {cnt:>3}")

    print()
    print("  ── תיקונים ──")
    for c in report["corrections"]:
        sev = SEVERITY_LEVELS[c["severity"]]
        print(f"  {sev['emoji']} #{c['id']} [{c['severity_he']}] {c['category_he']}")
        if c["draft_text"]:
            print(f"    − {c['draft_text'][:85]}")
        if c["final_text"]:
            print(f"    + {c['final_text'][:85]}")
        for wc in c.get("word_changes", [])[:3]:
            if wc["draft"] or wc["final"]:
                print(f"      [{wc['type']}] \"{wc['draft']}\" → \"{wc['final']}\"")
        a = c.get("analysis", {})
        if a.get("root_cause"):
            print(f"    💡 סיבה: {a['root_cause'][:90]}")
        if a.get("lesson"):
            print(f"    📘 לקח: {a['lesson'][:90]}")
        print()


# ---------------------------------------------------------------------------
# Demo
# ---------------------------------------------------------------------------

def _make_demo_docx(paragraphs, path):
    doc = Document()
    for t in paragraphs:
        doc.add_paragraph(t)
    doc.save(str(path))


DEMO_DRAFT = [
    "State of Israel",
    "Regional Rabbinical Court of Tel-Aviv-Jaffa",
    "Certificate of Divorce",
    "Case Number: 1234567/2025",
    "The Husband: Israel Ben David Kohen, Identity Card No. 012345678",
    "15 Hertzel Street, Tel Aviv-Jaffa",
    "The Wife: Sarah Bat Moshe Levy (nee Kohen), Identity Card No. 987654321",
    "42 Dizengof Street, Tel Aviv-Jaffa",
    "Were divorced according to Jewish Law on the 25th of Tevet 5786,",
    "corresponding to January 15, 2026, at the Regional Rabbinical Court of Tel-Aviv-Jaffa.",
    "The Get was delivered by the Husband to the Wife in the presence of the Court.",
    "From this date, both parties are free to marry any person, subject to the provisions of any law.",
    "Rabbi Abraham Yitzhaki — Head of Court",
    "Rabbi Moshe Davidov — Judge",
    "Rabbi Yaakov Friedman — Judge",
    "[Official Seal of the Regional Rabbinical Court of Tel-Aviv-Jaffa]",
    "Certificate No: GET/2026/4567",
]

DEMO_FINAL = [
    "State of Israel",
    "Regional Rabbinical Court of Tel Aviv-Yafo",
    "Certificate of Divorce",
    "Case Number: 1234567/2025",
    "The Husband: Israel Ben David Cohen, Identity Card No. 012345678",
    "15 Herzl Street, Tel Aviv-Yafo",
    "The Wife: Sarah Bat Moshe Levy (née Cohen), Identity Card No. 987654321",
    "42 Dizengoff Street, Tel Aviv-Yafo",
    "Were divorced in accordance with the Law of Moses and Israel on the 25th of Tevet 5786,",
    "corresponding to January 15, 2026, at the Regional Rabbinical Court of Tel Aviv-Yafo.",
    "The Get (bill of divorce) was delivered by the Husband to the Wife in the presence of the Court.",
    "From this date, both parties are free to marry any person, subject to the provisions of any applicable law.",
    "Rabbi Avraham Yitzhaki — Head of the Court (Av Beit Din)",
    "Rabbi Moshe Davidov — Judge (Dayan)",
    "Rabbi Yaakov Friedman — Judge (Dayan)",
    "[Official Seal of the Regional Rabbinical Court of Tel Aviv-Yafo]",
    "Certificate Number: GET/2026/4567",
]


def run_demo():
    demo_dir = OUTPUT_DIR / "_demo_learning"
    demo_dir.mkdir(parents=True, exist_ok=True)

    dp = demo_dir / "demo_draft.docx"
    fp = demo_dir / "demo_final.docx"
    _make_demo_docx(DEMO_DRAFT, dp)
    _make_demo_docx(DEMO_FINAL, fp)

    report = compare_documents(dp, fp)
    print_report(report)

    # Save to learned_corrections.json
    result = save_to_learned_corrections(report, "demo")
    print(f"  💾 learned_corrections.json: {result['new']} חדשים, {result['updated']} עודכנו, {result['total']} סה״כ")

    # Save report JSON
    rp = demo_dir / "demo_report.json"
    rp.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  💾 דו״ח: {rp}")

    # Show patterns
    db = _load_db()
    if db["patterns"]:
        print()
        print("  📖 Patterns שנשמרו (Agent 2 קורא מכאן):")
        for wrong, corrs in sorted(db["patterns"].items()):
            for c in corrs:
                print(f"    \"{wrong}\" → \"{c['correct']}\"  [{c['category']}, {c['severity']}]")
    print()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Learning Loop — השוואת DOCX + למידה")
    parser.add_argument("--draft", help="קובץ DOCX טיוטת AI")
    parser.add_argument("--final", help="קובץ DOCX גרסה סופית מאושרת")
    parser.add_argument("--out", help="שמור דו״ח JSON")
    parser.add_argument("--monday", action="store_true", help="דווח ל-monday.com")
    parser.add_argument("--demo", action="store_true", help="הרץ דוגמה מובנית")
    args = parser.parse_args()

    if args.demo:
        run_demo()
        return

    if not args.draft or not args.final:
        parser.error("נדרש --draft ו---final (או --demo)")

    report = compare_documents(Path(args.draft), Path(args.final))
    print_report(report)

    result = save_to_learned_corrections(report, str(args.draft))
    print(f"  💾 learned_corrections.json: {result['new']} חדשים, {result['total']} סה״כ")

    if args.out:
        Path(args.out).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"  💾 דו״ח: {args.out}")

    if args.monday:
        report_to_monday(report, result)


if __name__ == "__main__":
    main()
